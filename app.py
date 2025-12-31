
from __future__ import annotations

from flask import Flask, render_template, redirect, url_for, request, flash, jsonify, send_file
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy.orm import backref
from sqlalchemy import func, text
from flask_login import (
    LoginManager, login_user, logout_user,
    login_required, current_user, UserMixin
)
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import datetime, date, timedelta
from functools import wraps
import secrets
import re
import string
from collections import defaultdict
from io import BytesIO

# Word export
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ================== CẤU HÌNH APP ======================
app = Flask(__name__)
app.config["SECRET_KEY"] = "thay-chuoi-nay-bang-mot-secret-khac"
app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///lab_scheduler.db"
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False

db = SQLAlchemy(app)

login_manager = LoginManager(app)
login_manager.login_view = "login"


# ================== MODEL DỮ LIỆU ======================

class User(UserMixin, db.Model):
    __tablename__ = "user"
    id = db.Column(db.Integer, primary_key=True)
    full_name = db.Column(db.String(100), nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password_hash = db.Column(db.String(200), nullable=False)
    role = db.Column(db.String(20), nullable=False, default="teacher")  # admin/teacher

    group_name = db.Column(db.String(100), nullable=True)
    phone = db.Column(db.String(20), nullable=True)

    preferred_lab_id = db.Column(db.Integer, db.ForeignKey("lab.id"), nullable=True)
    preferred_lab = db.relationship(
        "Lab",
        foreign_keys=[preferred_lab_id],
        backref="preferred_teachers"
    )

    def set_password(self, password: str):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password: str) -> bool:
        """
        Tương thích DB cũ:
        - Nếu password_hash là dạng hash chuẩn -> check_password_hash
        - Nếu password_hash bị lưu dạng plain-text (cũ) -> cho đăng nhập rồi tự nâng cấp sang hash
        """
        if not self.password_hash:
            return False

        # Hash hiện đại thường bắt đầu bằng 'pbkdf2:' hoặc 'scrypt:'
        if self.password_hash.startswith(("pbkdf2:", "scrypt:")):
            return check_password_hash(self.password_hash, password)

        # fallback: plain text (legacy)
        if self.password_hash == password:
            # auto-upgrade
            self.set_password(password)
            db.session.commit()
            return True
        return False


class Lab(db.Model):
    __tablename__ = "lab"
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)

    # tổ/nhóm bộ môn (loại phòng)
    subject_group = db.Column(db.String(100), nullable=False, default="Tin học")

    capacity = db.Column(db.Integer, nullable=False, default=40)
    is_active = db.Column(db.Boolean, default=True)


class WeekSchedule(db.Model):
    __tablename__ = "week_schedule"
    id = db.Column(db.Integer, primary_key=True)

    school_year = db.Column(db.String(20), nullable=False, default="2025-2026")
    week_no = db.Column(db.Integer, nullable=False)

    start_date = db.Column(db.Date, nullable=False)
    end_date = db.Column(db.Date, nullable=False)

    is_active = db.Column(db.Boolean, default=True)

    __table_args__ = (
        db.UniqueConstraint("school_year", "week_no", name="uq_year_weekno"),
    )


class LabRequest(db.Model):
    __tablename__ = "lab_request"
    id = db.Column(db.Integer, primary_key=True)
    teacher_id = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=False)

    class_name = db.Column(db.String(100), nullable=False)
    num_students = db.Column(db.Integer, nullable=False, default=0)

    lab_group = db.Column(db.String(100), nullable=False, default="Tin học")

    week_id = db.Column(db.Integer, db.ForeignKey("week_schedule.id"), nullable=True)
    week = db.relationship("WeekSchedule", backref="lab_requests")

    date = db.Column(db.Date, nullable=False)
    period = db.Column(db.Integer, nullable=False)
    status = db.Column(db.String(20), nullable=False, default="pending")  # pending/approved/scheduled/conflict
    preferred_lab_id = db.Column(db.Integer, db.ForeignKey("lab.id"), nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

    teacher = db.relationship("User", backref="lab_requests")
    preferred_lab = db.relationship("Lab", foreign_keys=[preferred_lab_id])


class Reservation(db.Model):
    __tablename__ = "reservation"
    id = db.Column(db.Integer, primary_key=True)
    lab_id = db.Column(db.Integer, db.ForeignKey("lab.id"), nullable=False)
    request_id = db.Column(db.Integer, db.ForeignKey("lab_request.id"), nullable=False)
    date = db.Column(db.Date, nullable=False)
    period = db.Column(db.Integer, nullable=False)

    lab = db.relationship("Lab", backref="reservations")
    lab_request = db.relationship("LabRequest", backref=backref("reservation", uselist=False))


# ================== LOGIN MANAGER ======================

@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))


# ================== HÀM HỖ TRỢ ========================

def admin_required(func_):
    @wraps(func_)
    def wrapper(*args, **kwargs):
        if not current_user.is_authenticated or current_user.role != "admin":
            flash("Bạn không có quyền truy cập chức năng này.", "danger")
            return redirect(url_for("dashboard"))
        return func_(*args, **kwargs)
    return wrapper


def generate_password(length: int = 10) -> str:
    alphabet = string.ascii_letters + string.digits
    return "".join(secrets.choice(alphabet) for _ in range(length))


def get_month_range(year: int, month: int):
    start = date(year, month, 1)
    end = date(year + 1, 1, 1) if month == 12 else date(year, month + 1, 1)
    return start, end


def _week_date_from_weekday(week: WeekSchedule, weekday: int) -> date:
    return week.start_date + timedelta(days=weekday)


# ================== REPORT HELPERS ==================

def _get_stats_between(start_date: date, end_date_exclusive: date):
    teacher_stats = (
        db.session.query(
            User.full_name,
            User.group_name,
            User.phone,
            func.count(Reservation.id).label("total_periods"),
        )
        .join(LabRequest, LabRequest.teacher_id == User.id)
        .join(Reservation, Reservation.request_id == LabRequest.id)
        .filter(
            User.role == "teacher",
            Reservation.date >= start_date,
            Reservation.date < end_date_exclusive,
        )
        .group_by(User.full_name, User.group_name, User.phone)
        .order_by(User.full_name.asc())
        .all()
    )

    lab_stats = (
        db.session.query(
            Lab.name,
            Lab.subject_group,
            func.count(Reservation.id).label("total_periods"),
        )
        .join(Reservation, Reservation.lab_id == Lab.id)
        .filter(
            Reservation.date >= start_date,
            Reservation.date < end_date_exclusive,
        )
        .group_by(Lab.name, Lab.subject_group)
        .order_by(Lab.subject_group.asc(), Lab.name.asc())
        .all()
    )

    return teacher_stats, lab_stats


def _build_report_docx(title: str, start_date: date, end_date_inclusive: date, teacher_stats, lab_stats) -> BytesIO:
    doc = Document()

    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2 = doc.add_paragraph()
    run2 = p2.add_run(
        f"Khoảng thời gian: {start_date.strftime('%d/%m/%Y')} → {end_date_inclusive.strftime('%d/%m/%Y')}"
    )
    run2.font.size = Pt(11)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    doc.add_paragraph().add_run("Báo cáo theo giáo viên").bold = True
    t = doc.add_table(rows=1, cols=5)
    hdr = t.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Giáo viên"
    hdr[2].text = "Tổ/nhóm"
    hdr[3].text = "SĐT"
    hdr[4].text = "Số tiết thực hành"

    for i, row in enumerate(teacher_stats, start=1):
        r = t.add_row().cells
        r[0].text = str(i)
        r[1].text = row[0] or ""
        r[2].text = row[1] or ""
        r[3].text = row[2] or ""
        r[4].text = str(row[3] or 0)

    doc.add_paragraph("")

    doc.add_paragraph().add_run("Báo cáo tổng hợp theo phòng máy").bold = True
    t2 = doc.add_table(rows=1, cols=4)
    hdr2 = t2.rows[0].cells
    hdr2[0].text = "#"
    hdr2[1].text = "Phòng máy"
    hdr2[2].text = "Nhóm phòng"
    hdr2[3].text = "Số tiết sử dụng"

    for i, row in enumerate(lab_stats, start=1):
        r = t2.add_row().cells
        r[0].text = str(i)
        r[1].text = row[0] or ""
        r[2].text = row[1] or ""
        r[3].text = str(row[2] or 0)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# ================== SQLITE “TỰ NÂNG CẤP CỘT” ======================

def _sqlite_has_column(table_name: str, col_name: str) -> bool:
    rows = db.session.execute(text(f"PRAGMA table_info({table_name});")).fetchall()
    return any(r[1] == col_name for r in rows)


def ensure_schema_sqlite():
    try:
        if not _sqlite_has_column("lab", "subject_group"):
            db.session.execute(text(
                "ALTER TABLE lab ADD COLUMN subject_group VARCHAR(100) NOT NULL DEFAULT 'Tin học';"
            ))
            db.session.commit()

        if not _sqlite_has_column("lab_request", "lab_group"):
            db.session.execute(text(
                "ALTER TABLE lab_request ADD COLUMN lab_group VARCHAR(100) NOT NULL DEFAULT 'Tin học';"
            ))
            db.session.commit()

        if not _sqlite_has_column("lab_request", "week_id"):
            db.session.execute(text(
                "ALTER TABLE lab_request ADD COLUMN week_id INTEGER;"
            ))
            db.session.commit()

    except Exception as e:
        print("ensure_schema_sqlite ERROR:", e)


def ensure_default_labs_and_rename():
    mapping = {
        "Phòng máy 1": "Phòng thực hành Tin học 1",
        "Phòng máy 2": "Phòng thực hành Tin học 2",
        "Phòng máy 3": "Phòng thực hành Tin học 3",
    }

    for old_name, new_name in mapping.items():
        old_lab = Lab.query.filter_by(name=old_name).first()
        if old_lab:
            existed_new = Lab.query.filter_by(name=new_name).first()
            if existed_new:
                old_lab.is_active = False
            else:
                old_lab.name = new_name
            if not old_lab.subject_group:
                old_lab.subject_group = "Tin học"
    db.session.commit()

    desired = [
        ("Phòng thực hành Tin học 1", "Tin học", 40),
        ("Phòng thực hành Tin học 2", "Tin học", 40),
        ("Phòng thực hành Tin học 3", "Tin học", 36),
    ]
    for name, group, cap in desired:
        lab = Lab.query.filter_by(name=name).first()
        if not lab:
            db.session.add(Lab(name=name, subject_group=group, capacity=cap, is_active=True))
        else:
            if not lab.subject_group:
                lab.subject_group = group
    db.session.commit()


def init_default_users(force_reset_admin: bool = False):
    """
    Tạo admin mặc định (nếu chưa có).
    Nếu force_reset_admin=True -> reset password admin về admin123.
    """
    admin = User.query.filter(func.lower(User.email) == "admin@example.com").first()
    if not admin:
        admin = User(
            full_name="Quản trị hệ thống",
            email="admin@example.com",
            role="admin",
            group_name="Ban giám hiệu",
            phone="0123456789"
        )
        admin.set_password("admin123")
        db.session.add(admin)
        db.session.commit()
    else:
        if force_reset_admin:
            admin.set_password("admin123")
            db.session.commit()

    if not User.query.filter(func.lower(User.email) == "teacher1@example.com").first():
        lab1 = Lab.query.filter_by(name="Phòng thực hành Tin học 1").first()
        teacher = User(
            full_name="Giáo viên Tin học 1",
            email="teacher1@example.com",
            role="teacher",
            preferred_lab_id=lab1.id if lab1 else None,
            group_name="Tổ Tin học",
            phone="0987000000"
        )
        teacher.set_password("teacher123")
        db.session.add(teacher)
        db.session.commit()


def _infer_req_group(req: LabRequest) -> str:
    if getattr(req, "lab_group", None):
        return req.lab_group

    if req.preferred_lab_id:
        lab = db.session.get(Lab, req.preferred_lab_id)
        if lab and lab.subject_group:
            return lab.subject_group

    if req.teacher and req.teacher.preferred_lab_id:
        lab = db.session.get(Lab, req.teacher.preferred_lab_id)
        if lab and lab.subject_group:
            return lab.subject_group

    return "Tin học"


def _assign_reservation_for_request(req: LabRequest) -> bool:
    if Reservation.query.filter_by(request_id=req.id).first():
        return True

    req_group = _infer_req_group(req)
    req.lab_group = req_group
    db.session.flush()

    active_labs = Lab.query.filter_by(is_active=True, subject_group=req_group).order_by(Lab.name.asc()).all()
    if not active_labs:
        return False

    occupied = (
        db.session.query(Reservation.lab_id)
        .join(Lab, Reservation.lab_id == Lab.id)
        .filter(Reservation.date == req.date, Reservation.period == req.period, Lab.subject_group == req_group)
        .all()
    )
    occupied_ids = {lab_id for (lab_id,) in occupied}

    pref_lab_id = req.preferred_lab_id or (req.teacher.preferred_lab_id if req.teacher else None)
    chosen_lab = None

    if pref_lab_id:
        pref_lab = db.session.get(Lab, pref_lab_id)
        if pref_lab and pref_lab.is_active and pref_lab.subject_group == req_group and pref_lab.id not in occupied_ids:
            chosen_lab = pref_lab

    if not chosen_lab:
        for lab in active_labs:
            if lab.id not in occupied_ids:
                chosen_lab = lab
                break

    if not chosen_lab:
        return False

    db.session.add(Reservation(lab_id=chosen_lab.id, request_id=req.id, date=req.date, period=req.period))
    req.status = "scheduled"
    return True


def run_auto_schedule():
    unscheduled = LabRequest.query.filter(LabRequest.status.in_(["pending", "approved"])).all()
    unscheduled.sort(
        key=lambda r: (0 if (r.preferred_lab_id or (r.teacher and r.teacher.preferred_lab_id)) else 1, r.created_at)
    )

    for req in unscheduled:
        if Reservation.query.filter_by(request_id=req.id).first():
            continue

        ok = _assign_reservation_for_request(req)
        if not ok:
            req.status = "conflict"

    db.session.commit()


def ensure_reservations_for_scheduled_requests():
    scheduled_reqs = LabRequest.query.filter_by(status="scheduled").all()
    created = 0
    conflicted = 0
    for req in scheduled_reqs:
        if Reservation.query.filter_by(request_id=req.id).first():
            continue
        ok = _assign_reservation_for_request(req)
        if ok:
            created += 1
        else:
            req.status = "conflict"
            conflicted += 1
    db.session.commit()
    print(f"[SYNC] Backfill Reservation: created={created}, set_conflict={conflicted}")


def _get_occupied_details(target_date: date, target_period: int, subject_group: str):
    rows = (
        db.session.query(Reservation, Lab, LabRequest, User)
        .join(Lab, Reservation.lab_id == Lab.id)
        .join(LabRequest, Reservation.request_id == LabRequest.id)
        .join(User, LabRequest.teacher_id == User.id)
        .filter(Reservation.date == target_date, Reservation.period == target_period, Lab.subject_group == subject_group)
        .all()
    )
    out = []
    for res, lab, req, teacher in rows:
        out.append({
            "lab_name": lab.name,
            "teacher_name": teacher.full_name,
            "class_name": req.class_name
        })
    return out


# ================== AI: giải thích xung đột + gợi ý ==================

def ai_explain_conflict_and_suggest(req_id: int, max_period: int = 10, day_window: int = 7, top_k: int = 3):
    """
    Không dùng API AI bên ngoài để tránh lỗi.
    Gợi ý dựa trên phòng trống:
    - quét +/- day_window ngày quanh ngày yêu cầu
    - mỗi ngày: period 1..max_period
    - ưu tiên: cùng ngày, gần tiết, gần ngày
    """
    req = LabRequest.query.get(req_id)
    if not req:
        return {"ok": False, "message": "Không tìm thấy yêu cầu."}

    subject_group = _infer_req_group(req)
    active_labs = Lab.query.filter_by(is_active=True, subject_group=subject_group).order_by(Lab.name.asc()).all()

    occupied_now = _get_occupied_details(req.date, req.period, subject_group)
    reason = "Xung đột vì tất cả phòng phù hợp đã có lịch ở ngày/tiết này."

    suggestions = []
    # search around
    for delta_days in range(0, day_window + 1):
        for sign in ([0] if delta_days == 0 else [-1, 1]):
            d = req.date + timedelta(days=sign * delta_days)
            # skip weekend? giữ nguyên vì trường có thể học thứ 7
            for p in range(1, max_period + 1):
                if d == req.date and p == req.period:
                    continue
                # check if any lab free
                occupied = (
                    db.session.query(Reservation.lab_id)
                    .join(Lab, Reservation.lab_id == Lab.id)
                    .filter(Reservation.date == d, Reservation.period == p, Lab.subject_group == subject_group)
                    .all()
                )
                occupied_ids = {lab_id for (lab_id,) in occupied}
                free_labs = [lab for lab in active_labs if lab.id not in occupied_ids]
                if free_labs:
                    suggestions.append({
                        "date": d.isoformat(),
                        "date_str": d.strftime("%d/%m/%Y"),
                        "period": p,
                        "free_lab_names": [lab.name for lab in free_labs][:3],
                    })
                if len(suggestions) >= top_k:
                    break
            if len(suggestions) >= top_k:
                break
        if len(suggestions) >= top_k:
            break

    return {
        "ok": True,
        "req": {
            "id": req.id,
            "date": req.date.isoformat(),
            "date_str": req.date.strftime("%d/%m/%Y"),
            "period": req.period,
            "lab_group": subject_group,
            "class_name": req.class_name,
        },
        "reason": reason,
        "occupied": occupied_now,
        "suggestions": suggestions,
    }


# ================== ROUTES ===================

@app.route("/")
def index():
    return redirect(url_for("dashboard")) if current_user.is_authenticated else redirect(url_for("login"))


@app.route("/login", methods=["GET", "POST"])
def login():
    if current_user.is_authenticated:
        return redirect(url_for("dashboard"))

    if request.method == "POST":
        email = (request.form.get("email") or "").strip().lower()
        password = request.form.get("password") or ""

        user = User.query.filter(func.lower(User.email) == email).first()
        if user and user.check_password(password):
            login_user(user)
            flash("Đăng nhập thành công.", "success")
            return redirect(url_for("dashboard"))
        flash("Sai email hoặc mật khẩu.", "danger")

    return render_template("login.html")


@app.route("/logout")
@login_required
def logout():
    logout_user()
    flash("Đã đăng xuất.", "info")
    return redirect(url_for("login"))


@app.route("/change_password", methods=["GET", "POST"])
@login_required
def change_password():
    if request.method == "POST":
        current_pw = request.form.get("current_password", "")
        new_pw = request.form.get("new_password", "")
        confirm_pw = request.form.get("confirm_password", "")

        if not current_pw or not new_pw or not confirm_pw:
            flash("Vui lòng nhập đầy đủ thông tin.", "danger")
            return render_template("change_password.html")

        if not current_user.check_password(current_pw):
            flash("Mật khẩu hiện tại không đúng.", "danger")
            return render_template("change_password.html")

        if new_pw != confirm_pw:
            flash("Mật khẩu mới nhập lại không khớp.", "danger")
            return render_template("change_password.html")

        current_user.set_password(new_pw)
        db.session.commit()
        flash("Đã đổi mật khẩu thành công.", "success")
        return redirect(url_for("dashboard"))

    return render_template("change_password.html")


@app.route("/dashboard")
@login_required
def dashboard():
    if current_user.role == "admin":
        total_requests = LabRequest.query.count()
        scheduled_count = LabRequest.query.filter_by(status="scheduled").count()
        conflict_count = LabRequest.query.filter_by(status="conflict").count()
        pending_count = LabRequest.query.filter(LabRequest.status.in_(["pending", "approved"])).count()

        return render_template(
            "dashboard_admin.html",
            total_requests=total_requests,
            scheduled_count=scheduled_count,
            conflict_count=conflict_count,
            pending_count=pending_count
        )

    my_requests = LabRequest.query.filter_by(teacher_id=current_user.id).order_by(
        LabRequest.date.desc(), LabRequest.period.desc()
    ).all()
    return render_template("dashboard_teacher.html", my_requests=my_requests)


# ================== ADMIN: TUẦN ĐĂNG KÝ ==================

@app.route("/admin/weeks", methods=["GET", "POST"])
@login_required
@admin_required
def admin_weeks():
    if request.method == "POST":
        school_year = (request.form.get("school_year") or "2025-2026").strip()
        week_no_str = (request.form.get("week_no") or "").strip()
        start_str = (request.form.get("start_date") or "").strip()
        end_str = (request.form.get("end_date") or "").strip()

        if not week_no_str or not start_str or not end_str:
            flash("Vui lòng nhập đủ: Năm học, Tuần, Từ ngày, Đến ngày.", "danger")
            return redirect(url_for("admin_weeks"))

        try:
            week_no = int(week_no_str)
        except ValueError:
            flash("Tuần phải là số.", "danger")
            return redirect(url_for("admin_weeks"))

        try:
            start_date = datetime.strptime(start_str, "%Y-%m-%d").date()
            end_date = datetime.strptime(end_str, "%Y-%m-%d").date()
        except ValueError:
            flash("Ngày không hợp lệ (YYYY-MM-DD).", "danger")
            return redirect(url_for("admin_weeks"))

        if end_date < start_date:
            flash("Đến ngày phải >= Từ ngày.", "danger")
            return redirect(url_for("admin_weeks"))

        existed = WeekSchedule.query.filter_by(school_year=school_year, week_no=week_no).first()
        if existed:
            existed.start_date = start_date
            existed.end_date = end_date
            existed.is_active = True
        else:
            db.session.add(WeekSchedule(
                school_year=school_year,
                week_no=week_no,
                start_date=start_date,
                end_date=end_date,
                is_active=True
            ))

        db.session.commit()
        flash("Đã lưu tuần học.", "success")
        return redirect(url_for("admin_weeks"))

    weeks = WeekSchedule.query.order_by(WeekSchedule.school_year.desc(), WeekSchedule.week_no.asc()).all()
    return render_template("admin_weeks.html", weeks=weeks)


@app.route("/admin/weeks/<int:week_id>/toggle", methods=["POST"])
@login_required
@admin_required
def admin_toggle_week(week_id):
    w = WeekSchedule.query.get_or_404(week_id)
    w.is_active = not w.is_active
    db.session.commit()
    flash("Đã cập nhật trạng thái tuần.", "success")
    return redirect(url_for("admin_weeks"))


@app.route("/admin/weeks/<int:week_id>/update", methods=["POST"])
@login_required
@admin_required
def admin_update_week(week_id):
    w = WeekSchedule.query.get_or_404(week_id)

    school_year = (request.form.get("school_year") or w.school_year).strip()
    week_no_str = (request.form.get("week_no") or str(w.week_no)).strip()
    start_str = (request.form.get("start_date") or "").strip()
    end_str = (request.form.get("end_date") or "").strip()
    is_active = True if (request.form.get("is_active") == "1") else False

    try:
        week_no = int(week_no_str)
    except ValueError:
        flash("Tuần phải là số.", "danger")
        return redirect(url_for("admin_weeks"))

    try:
        start_date = datetime.strptime(start_str, "%Y-%m-%d").date()
        end_date = datetime.strptime(end_str, "%Y-%m-%d").date()
    except ValueError:
        flash("Ngày không hợp lệ (YYYY-MM-DD).", "danger")
        return redirect(url_for("admin_weeks"))

    if end_date < start_date:
        flash("Đến ngày phải >= Từ ngày.", "danger")
        return redirect(url_for("admin_weeks"))

    conflict = WeekSchedule.query.filter(
        WeekSchedule.id != w.id,
        WeekSchedule.school_year == school_year,
        WeekSchedule.week_no == week_no
    ).first()
    if conflict:
        flash("Đã tồn tại tuần này trong cùng năm học.", "danger")
        return redirect(url_for("admin_weeks"))

    w.school_year = school_year
    w.week_no = week_no
    w.start_date = start_date
    w.end_date = end_date
    w.is_active = is_active

    db.session.commit()
    flash("Đã cập nhật tuần.", "success")
    return redirect(url_for("admin_weeks"))


@app.route("/admin/weeks/<int:week_id>/delete", methods=["POST"])
@login_required
@admin_required
def admin_delete_week(week_id):
    w = WeekSchedule.query.get_or_404(week_id)

    used = LabRequest.query.filter_by(week_id=w.id).first()
    if used:
        w.is_active = False
        db.session.commit()
        flash("Tuần đã có giáo viên đăng ký. Hệ thống đã TẮT tuần thay vì xoá.", "warning")
        return redirect(url_for("admin_weeks"))

    db.session.delete(w)
    db.session.commit()
    flash("Đã xoá tuần.", "success")
    return redirect(url_for("admin_weeks"))


# ================== ADMIN: YÊU CẦU & XẾP LỊCH ==================

@app.route("/admin/requests")
@login_required
@admin_required
def admin_requests():
    all_requests = LabRequest.query.order_by(LabRequest.date.desc(), LabRequest.period.desc()).all()
    labs = Lab.query.filter_by(is_active=True).order_by(Lab.subject_group.asc(), Lab.name.asc()).all()

    conflict_map = {}
    for r in all_requests:
        if r.status == "conflict":
            req_group = _infer_req_group(r)
            conflict_map[r.id] = _get_occupied_details(r.date, r.period, req_group)

    return render_template(
        "admin_requests.html",
        all_requests=all_requests,
        labs=labs,
        conflict_map=conflict_map
    )


@app.route("/admin/request/<int:req_id>/update_status", methods=["POST"])
@login_required
@admin_required
def update_request_status(req_id):
    lab_request = LabRequest.query.get_or_404(req_id)
    new_status = (request.form.get("status") or "").strip()

    if new_status not in ["pending", "approved", "scheduled", "conflict"]:
        flash("Trạng thái không hợp lệ.", "danger")
        return redirect(url_for("admin_requests"))

    if new_status == "scheduled":
        ok = _assign_reservation_for_request(lab_request)
        if not ok:
            lab_request.status = "conflict"
            db.session.commit()
            flash("Không thể xếp lịch vì không còn phòng trống. Đã chuyển 'Xung đột'.", "danger")
            return redirect(url_for("admin_requests"))

        db.session.commit()
        flash("Đã xếp lịch và cập nhật trạng thái.", "success")
        return redirect(url_for("admin_requests"))

    if lab_request.reservation:
        db.session.delete(lab_request.reservation)

    lab_request.status = new_status
    db.session.commit()
    flash("Đã cập nhật trạng thái.", "success")
    return redirect(url_for("admin_requests"))


@app.route("/admin/request/<int:req_id>/manual_assign", methods=["POST"])
@login_required
@admin_required
def manual_assign(req_id):
    lab_request = LabRequest.query.get_or_404(req_id)
    lab_id_str = (request.form.get("lab_id") or "").strip()

    if not lab_id_str:
        flash("Vui lòng chọn phòng khi gán thủ công.", "danger")
        return redirect(url_for("admin_requests"))

    try:
        lab_id = int(lab_id_str)
    except ValueError:
        flash("Phòng không hợp lệ.", "danger")
        return redirect(url_for("admin_requests"))

    lab = Lab.query.get(lab_id)
    if not lab or not lab.is_active:
        flash("Phòng không tồn tại hoặc không hoạt động.", "danger")
        return redirect(url_for("admin_requests"))

    req_group = _infer_req_group(lab_request)
    if lab.subject_group != req_group:
        flash(f"Không thể gán. Yêu cầu nhóm '{req_group}' nhưng phòng thuộc '{lab.subject_group}'.", "danger")
        return redirect(url_for("admin_requests"))

    conflict = Reservation.query.filter_by(lab_id=lab.id, date=lab_request.date, period=lab_request.period).first()
    if conflict:
        flash("Phòng này đã có lịch ở ngày/tiết này.", "danger")
        return redirect(url_for("admin_requests"))

    if lab_request.reservation:
        db.session.delete(lab_request.reservation)

    db.session.add(Reservation(lab_id=lab.id, request_id=lab_request.id, date=lab_request.date, period=lab_request.period))
    lab_request.status = "scheduled"
    db.session.commit()

    flash("Đã gán phòng thủ công.", "success")
    return redirect(url_for("admin_requests"))


@app.route("/admin/auto_schedule", methods=["POST"])
@login_required
@admin_required
def auto_schedule():
    run_auto_schedule()
    flash("Đã chạy xếp lịch tự động.", "success")
    return redirect(url_for("admin_requests"))


# ================== API: AI GỢI Ý ==================

@app.route("/api/requests/<int:req_id>/ai_conflict", methods=["GET"])
@login_required
def api_ai_conflict(req_id):
    req = LabRequest.query.get_or_404(req_id)
    if current_user.role != "admin" and req.teacher_id != current_user.id:
        return jsonify({"error": "Bạn không có quyền xem yêu cầu này."}), 403

    data = ai_explain_conflict_and_suggest(req_id=req_id, max_period=10, day_window=7, top_k=3)
    return jsonify(data)


# ✅ Alias để tránh lỗi 404 do JS cũ gọi /api/request/<id>/ai_conflict
@app.route("/api/request/<int:req_id>/ai_conflict", methods=["GET"])
@login_required
def api_ai_conflict_alias(req_id):
    return api_ai_conflict(req_id)


# ================== GV: ĐĂNG KÝ ==================

@app.route("/request/new", methods=["GET", "POST"])
@login_required
def new_request():
    if current_user.role != "teacher":
        flash("Chỉ giáo viên mới được đăng ký phòng.", "danger")
        return redirect(url_for("dashboard"))

    active_labs = Lab.query.filter_by(is_active=True).order_by(Lab.subject_group.asc(), Lab.name.asc()).all()
    groups = sorted({lab.subject_group for lab in active_labs}) or ["Tin học"]
    weeks = WeekSchedule.query.filter_by(is_active=True).order_by(
        WeekSchedule.school_year.desc(), WeekSchedule.week_no.desc()
    ).all()

    if request.method == "POST":
        class_name = (request.form.get("class_name") or "").strip()
        period_str = (request.form.get("period") or "").strip()
        lab_group = (request.form.get("lab_group") or "").strip() or "Tin học"
        preferred_lab_id_str = (request.form.get("preferred_lab_id") or "").strip()

        week_id_str = (request.form.get("week_id") or "").strip()
        weekday_str = (request.form.get("weekday") or "").strip()
        date_str = (request.form.get("date") or "").strip()

        if not class_name or not period_str:
            flash("Vui lòng nhập đầy đủ thông tin.", "danger")
            return render_template("request_form.html", labs=active_labs, groups=groups, weeks=weeks)

        try:
            period = int(period_str)
        except ValueError:
            flash("Tiết không hợp lệ.", "danger")
            return render_template("request_form.html", labs=active_labs, groups=groups, weeks=weeks)

        date_obj = None
        week_obj = None

        if week_id_str:
            try:
                week_id = int(week_id_str)
            except ValueError:
                flash("Tuần không hợp lệ.", "danger")
                return render_template("request_form.html", labs=active_labs, groups=groups, weeks=weeks)

            week_obj = WeekSchedule.query.get(week_id)
            if not week_obj or not week_obj.is_active:
                flash("Tuần không tồn tại hoặc không hoạt động.", "danger")
                return render_template("request_form.html", labs=active_labs, groups=groups, weeks=weeks)

            if weekday_str == "":
                flash("Vui lòng chọn Thứ trong tuần.", "danger")
                return render_template("request_form.html", labs=active_labs, groups=groups, weeks=weeks)

            try:
                weekday = int(weekday_str)
            except ValueError:
                flash("Thứ không hợp lệ.", "danger")
                return render_template("request_form.html", labs=active_labs, groups=groups, weeks=weeks)

            if weekday < 0 or weekday > 6:
                flash("Thứ phải trong khoảng 0..6 (0=Thứ 2).", "danger")
                return render_template("request_form.html", labs=active_labs, groups=groups, weeks=weeks)

            date_obj = _week_date_from_weekday(week_obj, weekday)
            if date_obj < week_obj.start_date or date_obj > week_obj.end_date:
                flash("Ngày suy ra không nằm trong khoảng tuần.", "danger")
                return render_template("request_form.html", labs=active_labs, groups=groups, weeks=weeks)
        else:
            if not date_str:
                flash("Vui lòng chọn ngày đăng ký.", "danger")
                return render_template("request_form.html", labs=active_labs, groups=groups, weeks=weeks)
            try:
                date_obj = datetime.strptime(date_str, "%Y-%m-%d").date()
            except ValueError:
                flash("Ngày không hợp lệ.", "danger")
                return render_template("request_form.html", labs=active_labs, groups=groups, weeks=weeks)

        preferred_lab_id = None
        if preferred_lab_id_str:
            try:
                preferred_lab_id = int(preferred_lab_id_str)
            except ValueError:
                preferred_lab_id = None

        if preferred_lab_id:
            lab = db.session.get(Lab, preferred_lab_id)
            if not lab or not lab.is_active:
                flash("Phòng ưu tiên không hợp lệ.", "danger")
                return render_template("request_form.html", labs=active_labs, groups=groups, weeks=weeks)
            lab_group = lab.subject_group

        existed_same = (
            LabRequest.query
            .filter_by(teacher_id=current_user.id, date=date_obj, period=period)
            .filter(LabRequest.status.in_(["pending", "approved", "scheduled"]))
            .first()
        )
        if existed_same:
            flash("Bạn đã có yêu cầu/lịch ở ngày & tiết này. Hãy chọn tiết khác.", "danger")
            return render_template("request_form.html", labs=active_labs, groups=groups, weeks=weeks)

        lab_request = LabRequest(
            teacher_id=current_user.id,
            class_name=class_name,
            num_students=0,
            lab_group=lab_group,
            week_id=(week_obj.id if week_obj else None),
            date=date_obj,
            period=period,
            preferred_lab_id=preferred_lab_id,
            status="pending"
        )
        db.session.add(lab_request)
        db.session.commit()

        flash("Đã gửi yêu cầu đăng ký phòng.", "success")
        return redirect(url_for("dashboard"))

    return render_template("request_form.html", labs=active_labs, groups=groups, weeks=weeks)


# ================== XEM LỊCH ==================

@app.route("/schedule")
@login_required
def schedule():
    view_by = (request.args.get("by") or "").strip().lower()

    rows = (
        db.session.query(Reservation, Lab, LabRequest, User)
        .join(Lab, Reservation.lab_id == Lab.id)
        .join(LabRequest, Reservation.request_id == LabRequest.id)
        .join(User, LabRequest.teacher_id == User.id)
        .order_by(User.full_name.asc(), Reservation.date.asc(), Reservation.period.asc(), Lab.name.asc())
        .all()
    )

    if current_user.role == "teacher":
        rows = [t for t in rows if t[3].id == current_user.id]

    if view_by == "teacher":
        blocks_map = {}
        for res, lab, req, teacher in rows:
            item = {
                "date": res.date,
                "date_str": res.date.strftime("%d/%m/%Y"),
                "period": res.period,
                "lab_name": lab.name,
                "subject_group": lab.subject_group,
                "class_name": req.class_name
            }
            if teacher.id not in blocks_map:
                blocks_map[teacher.id] = {
                    "teacher_id": teacher.id,
                    "teacher_name": teacher.full_name,
                    "group_name": teacher.group_name,
                    "rows": []
                }
            blocks_map[teacher.id]["rows"].append(item)

        teacher_blocks = list(blocks_map.values())
        teacher_blocks.sort(key=lambda b: (b["teacher_name"] or "").lower())
        for b in teacher_blocks:
            b["rows"].sort(key=lambda x: (x["date"], x["period"], x["lab_name"]))

        return render_template("schedule_by_teacher.html", teacher_blocks=teacher_blocks)

    # lịch chung
    if current_user.role == "teacher":
        reservations = (
            Reservation.query
            .join(LabRequest, Reservation.request_id == LabRequest.id)
            .filter(LabRequest.teacher_id == current_user.id)
            .order_by(Reservation.date.asc(), Reservation.period.asc(), Reservation.lab_id.asc())
            .all()
        )
    else:
        reservations = Reservation.query.order_by(
            Reservation.date.asc(), Reservation.period.asc(), Reservation.lab_id.asc()
        ).all()

    return render_template("schedule.html", reservations=reservations)


# Alias endpoint cũ: url_for('view_schedule')
app.add_url_rule("/schedule", endpoint="view_schedule", view_func=schedule)


# ================== ADMIN: PHÒNG THỰC HÀNH ==================

@app.route("/admin/labs", methods=["GET", "POST"])
@login_required
@admin_required
def admin_labs():
    if request.method == "POST":
        name = (request.form.get("name") or "").strip()
        subject_group = (request.form.get("subject_group") or "").strip()

        if not name or not subject_group:
            flash("Vui lòng nhập Tên phòng và Tổ/Nhóm bộ môn.", "danger")
            return redirect(url_for("admin_labs"))

        if Lab.query.filter_by(name=name).first():
            flash("Tên phòng đã tồn tại.", "danger")
            return redirect(url_for("admin_labs"))

        db.session.add(Lab(name=name, subject_group=subject_group, capacity=40, is_active=True))
        db.session.commit()

        flash("Đã thêm phòng thực hành.", "success")
        return redirect(url_for("admin_labs"))

    labs = Lab.query.order_by(Lab.subject_group.asc(), Lab.name.asc()).all()
    groups = sorted({lab.subject_group for lab in labs}) or ["Tin học"]
    return render_template("admin_labs.html", labs=labs, groups=groups)


@app.route("/admin/labs/<int:lab_id>/toggle", methods=["POST"])
@login_required
@admin_required
def admin_toggle_lab(lab_id):
    lab = Lab.query.get_or_404(lab_id)
    lab.is_active = not lab.is_active
    db.session.commit()
    flash("Đã cập nhật trạng thái phòng.", "success")
    return redirect(url_for("admin_labs"))


@app.route("/admin/labs/<int:lab_id>/delete", methods=["POST"])
@login_required
@admin_required
def admin_delete_lab(lab_id):
    lab = Lab.query.get_or_404(lab_id)

    has_reservation = Reservation.query.filter_by(lab_id=lab.id).first()
    if has_reservation:
        flash("Không thể xoá vì phòng này đã có lịch xếp. Hãy 'Tắt phòng' thay vì xoá.", "danger")
        return redirect(url_for("admin_labs"))

    used_by_teacher = User.query.filter_by(preferred_lab_id=lab.id).first()
    if used_by_teacher:
        flash("Không thể xoá vì có giáo viên đang đặt phòng này làm phòng ưu tiên.", "danger")
        return redirect(url_for("admin_labs"))

    used_in_requests = LabRequest.query.filter_by(preferred_lab_id=lab.id).first()
    if used_in_requests:
        flash("Không thể xoá vì có yêu cầu đăng ký đang chọn phòng này.", "danger")
        return redirect(url_for("admin_labs"))

    db.session.delete(lab)
    db.session.commit()
    flash("Đã xoá phòng thực hành.", "success")
    return redirect(url_for("admin_labs"))


# ================== ADMIN: GIÁO VIÊN ==================

@app.route("/admin/teachers")
@login_required
@admin_required
def admin_teachers():
    teachers = User.query.filter_by(role="teacher").order_by(User.full_name).all()
    labs = Lab.query.filter_by(is_active=True).order_by(Lab.subject_group.asc(), Lab.name.asc()).all()
    return render_template("admin_teachers.html", teachers=teachers, labs=labs)


@app.route("/admin/teachers/new", methods=["GET", "POST"])
@login_required
@admin_required
def admin_new_teacher():
    labs = Lab.query.filter_by(is_active=True).order_by(Lab.subject_group.asc(), Lab.name.asc()).all()

    if request.method == "POST":
        full_name = (request.form.get("full_name") or "").strip()
        email = (request.form.get("email") or "").strip().lower()
        password = (request.form.get("password") or "").strip()
        group_name = (request.form.get("group_name") or "").strip()
        phone = (request.form.get("phone") or "").strip()
        preferred_lab_id_str = (request.form.get("preferred_lab_id") or "").strip()

        if not full_name or not email or not password:
            flash("Họ tên, email, mật khẩu là bắt buộc.", "danger")
            return render_template("teacher_form.html", labs=labs)

        if User.query.filter(func.lower(User.email) == email).first():
            flash("Email này đã tồn tại.", "danger")
            return render_template("teacher_form.html", labs=labs)

        preferred_lab_id = None
        if preferred_lab_id_str:
            try:
                preferred_lab_id = int(preferred_lab_id_str)
            except ValueError:
                preferred_lab_id = None

        teacher = User(
            full_name=full_name,
            email=email,
            role="teacher",
            group_name=group_name or None,
            phone=phone or None,
            preferred_lab_id=preferred_lab_id
        )
        teacher.set_password(password)
        db.session.add(teacher)
        db.session.commit()

        flash("Đã tạo tài khoản giáo viên mới.", "success")
        return redirect(url_for("admin_teachers"))

    return render_template("teacher_form.html", labs=labs)


@app.route("/admin/teachers/<int:teacher_id>/delete", methods=["POST"])
@login_required
@admin_required
def admin_delete_teacher(teacher_id):
    teacher = User.query.get_or_404(teacher_id)

    if teacher.role != "teacher":
        flash("Không thể xoá tài khoản không phải giáo viên.", "danger")
        return redirect(url_for("admin_teachers"))

    if teacher.id == current_user.id:
        flash("Không thể tự xoá tài khoản đang đăng nhập.", "danger")
        return redirect(url_for("admin_teachers"))

    for req in list(teacher.lab_requests):
        if req.reservation:
            db.session.delete(req.reservation)
        db.session.delete(req)

    db.session.delete(teacher)
    db.session.commit()

    flash("Đã xoá giáo viên + toàn bộ yêu cầu/lịch.", "success")
    return redirect(url_for("admin_teachers"))


@app.route("/admin/teachers/<int:teacher_id>/reset_password", methods=["POST"])
@login_required
@admin_required
def admin_reset_teacher_password(teacher_id):
    teacher = User.query.get_or_404(teacher_id)
    if teacher.role != "teacher":
        flash("Chỉ reset mật khẩu cho giáo viên.", "danger")
        return redirect(url_for("admin_teachers"))

    new_pw = (request.form.get("new_password") or "").strip()
    if not new_pw:
        new_pw = generate_password(10)

    teacher.set_password(new_pw)
    db.session.commit()
    flash(f"Đã reset MK cho {teacher.full_name}. Mật khẩu mới: {new_pw}", "warning")
    return redirect(url_for("admin_teachers"))


# ================== ADMIN: BÁO CÁO (TUẦN / KHOẢNG / THÁNG + WORD) ==================

@app.route("/admin/reports", methods=["GET", "POST"])
@login_required
@admin_required
def admin_reports():
    weeks = WeekSchedule.query.order_by(WeekSchedule.school_year.desc(), WeekSchedule.week_no.desc()).all()
    today = date.today()
    report_data = None

    if request.method == "POST":
        mode = (request.form.get("mode") or "").strip().lower()
        export = (request.form.get("export") or "").strip() == "1"

        start_date = None
        end_date_inclusive = None

        if mode == "week":
            week_id_str = (request.form.get("week_id") or "").strip()
            if not week_id_str:
                flash("Vui lòng chọn tuần để xem báo cáo.", "danger")
                return render_template("admin_reports.html", weeks=weeks, report_data=None, today=today)

            try:
                week_id = int(week_id_str)
            except ValueError:
                flash("Tuần không hợp lệ.", "danger")
                return render_template("admin_reports.html", weeks=weeks, report_data=None, today=today)

            w = WeekSchedule.query.get(week_id)
            if not w:
                flash("Tuần không tồn tại.", "danger")
                return render_template("admin_reports.html", weeks=weeks, report_data=None, today=today)

            start_date = w.start_date
            end_date_inclusive = w.end_date
            title = f"BÁO CÁO SỬ DỤNG PHÒNG THỰC HÀNH - Tuần {w.week_no} ({w.school_year})"

        elif mode == "range":
            start_str = (request.form.get("start_date") or "").strip()
            end_str = (request.form.get("end_date") or "").strip()

            if not start_str or not end_str:
                flash("Vui lòng nhập đủ Từ ngày / Đến ngày.", "danger")
                return render_template("admin_reports.html", weeks=weeks, report_data=None, today=today)

            try:
                start_date = datetime.strptime(start_str, "%Y-%m-%d").date()
                end_date_inclusive = datetime.strptime(end_str, "%Y-%m-%d").date()
            except ValueError:
                flash("Ngày không hợp lệ (định dạng YYYY-MM-DD).", "danger")
                return render_template("admin_reports.html", weeks=weeks, report_data=None, today=today)

            if end_date_inclusive < start_date:
                flash("Đến ngày phải >= Từ ngày.", "danger")
                return render_template("admin_reports.html", weeks=weeks, report_data=None, today=today)

            title = "BÁO CÁO SỬ DỤNG PHÒNG THỰC HÀNH - Theo khoảng thời gian"

        elif mode == "month":
            month_str = (request.form.get("month") or "").strip()
            year_str = (request.form.get("year") or "").strip()

            if not month_str or not year_str:
                flash("Vui lòng chọn Tháng và Năm.", "danger")
                return render_template("admin_reports.html", weeks=weeks, report_data=None, today=today)

            try:
                month = int(month_str)
                year = int(year_str)
            except ValueError:
                flash("Tháng/Năm không hợp lệ.", "danger")
                return render_template("admin_reports.html", weeks=weeks, report_data=None, today=today)

            if month < 1 or month > 12:
                flash("Tháng phải trong khoảng 1..12.", "danger")
                return render_template("admin_reports.html", weeks=weeks, report_data=None, today=today)

            start_date, end_excl = get_month_range(year, month)
            end_date_inclusive = end_excl - timedelta(days=1)
            title = f"BÁO CÁO SỬ DỤNG PHÒNG THỰC HÀNH - Tháng {month}/{year}"

        else:
            flash("Chế độ báo cáo không hợp lệ.", "danger")
            return render_template("admin_reports.html", weeks=weeks, report_data=None, today=today)

        end_exclusive = end_date_inclusive + timedelta(days=1)
        teacher_stats, lab_stats = _get_stats_between(start_date, end_exclusive)

        if export:
            bio = _build_report_docx(title, start_date, end_date_inclusive, teacher_stats, lab_stats)
            safe_title = re.sub(r"[^0-9A-Za-zÀ-ỹ _-]+", "", title).strip().replace(" ", "_")
            filename = f"{safe_title}_{start_date.strftime('%Y%m%d')}-{end_date_inclusive.strftime('%Y%m%d')}.docx"
            return send_file(
                bio,
                as_attachment=True,
                download_name=filename,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        report_data = {
            "mode": mode,
            "title": title,
            "start_date": start_date,
            "end_date": end_date_inclusive,
            "teacher_stats": teacher_stats,
            "lab_stats": lab_stats,
            "form": dict(request.form),
        }

    return render_template("admin_reports.html", weeks=weeks, report_data=report_data, today=today)


@app.route("/admin/reports/monthly")
@login_required
@admin_required
def admin_monthly_report():
    return redirect(url_for("admin_reports"))

def setup_app():
    """Khởi tạo DB và dữ liệu mặc định, dùng cho cả local & Render."""
    with app.app_context():
        db.create_all()
        ensure_schema_sqlite()
        ensure_default_labs_and_rename()
        init_default_users()
        ensure_reservations_for_scheduled_requests()

# GỌI NGAY KHI IMPORT MODULE (Render cũng chạy đoạn này)
setup_app()

# ================== MAIN ===========================
if __name__ == "__main__":
    with app.app_context():
        db.create_all()
        ensure_schema_sqlite()
        ensure_default_labs_and_rename()
        init_default_users(force_reset_admin=False)
        ensure_reservations_for_scheduled_requests()

    app.run(host="127.0.0.1", port=5000, debug=True, use_reloader=False)
